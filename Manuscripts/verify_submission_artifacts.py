"""Preflight checks for the active NEJM AI manuscript package.

The current manuscript workflow is direct DOCX maintenance. This verifier
checks the active manuscript and supplement plus the current onset/offset
survival artifacts. It intentionally does not require the archived title-page,
cover-letter, or manifest files from the retired generator workflow.
"""

from __future__ import annotations

import json
import re
import sys
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "Manuscripts" / "generated"
TABLE_DIR = OUT / "tables"
FIGURE_DIR = OUT / "figures"
ANALYSIS = ROOT / "Data" / "analysis"

MAIN_DOCX = OUT / "NEJM_AI_manuscript 7.2.2026.docx"
SUPPLEMENT_DOCX = OUT / "NEJM_AI_supplement 7.2.2026.docx"

REQUIRED_FILES = [
    MAIN_DOCX,
    SUPPLEMENT_DOCX,
    TABLE_DIR / "table1_variables_by_domain.csv",
    TABLE_DIR / "table2_onset_performance.csv",
    TABLE_DIR / "table3_offset_performance.csv",
    TABLE_DIR / "tableA1_generic_vs_brand.csv",
    TABLE_DIR / "tableA2_top_therapeutic_classes.csv",
    TABLE_DIR / "tableA3_logistic_feature_sensitivity.csv",
    TABLE_DIR / "tableA3_top20_onset_risk.csv",
    TABLE_DIR / "tableA4_top20_offset_resolve.csv",
    TABLE_DIR / "tableA5_feature_missingness.csv",
    TABLE_DIR / "tableA5_feature_missingness.xlsx",
    TABLE_DIR / "tableA8_feature_missingness_summary.csv",
    TABLE_DIR / "tableA11_no_prescription_benchmark.csv",
    TABLE_DIR / "tableA12_operational_yield.csv",
    TABLE_DIR / "tableA13_posthoc_calibration.csv",
    TABLE_DIR / "tableA14_horizon6_onset.csv",
    TABLE_DIR / "reference_audit.csv",
    TABLE_DIR / "reporting_checklist_author_review.csv",
    TABLE_DIR / "submission_admin_fields.csv",
    TABLE_DIR / "submission_readiness_checklist.csv",
    TABLE_DIR / "suggested_data_use_language.csv",
    TABLE_DIR / "suggested_disclosure_language.csv",
    FIGURE_DIR / "fig1_onset_shap.png",
    FIGURE_DIR / "fig1_onset_shap.tiff",
    FIGURE_DIR / "fig1_onset_shap.pdf",
    FIGURE_DIR / "fig1_onset_shap.svg",
    FIGURE_DIR / "fig1_onset_shap_full25.png",
    FIGURE_DIR / "fig1_onset_shap_full25.tiff",
    FIGURE_DIR / "fig1_onset_shap_full25.pdf",
    FIGURE_DIR / "fig1_onset_shap_full25.svg",
    FIGURE_DIR / "fig2_offset_shap.png",
    FIGURE_DIR / "fig2_offset_shap.tiff",
    FIGURE_DIR / "fig2_offset_shap.pdf",
    FIGURE_DIR / "fig2_offset_shap.svg",
    FIGURE_DIR / "fig2_offset_shap_full25.png",
    FIGURE_DIR / "fig2_offset_shap_full25.tiff",
    FIGURE_DIR / "fig2_offset_shap_full25.pdf",
    FIGURE_DIR / "fig2_offset_shap_full25.svg",
    FIGURE_DIR / "figA1_event_frequency.png",
    FIGURE_DIR / "figA1_event_frequency.tiff",
    FIGURE_DIR / "figA1_event_frequency.pdf",
    FIGURE_DIR / "figA1_event_frequency.svg",
    FIGURE_DIR / "figA2_calibration.png",
    FIGURE_DIR / "figA2_calibration.tiff",
    FIGURE_DIR / "figA2_calibration.pdf",
    FIGURE_DIR / "figA2_calibration.svg",
    ANALYSIS / "survival_onset" / "pooled_metrics.csv",
    ANALYSIS / "survival_onset" / "bootstrap_ci.csv",
    ANALYSIS / "survival_onset" / "harrell_c.csv",
    ANALYSIS / "survival_onset" / "permutation_importance.csv",
    ANALYSIS / "onset_without_licensed_prescription_full" / "survival_onset" / "pooled_metrics.csv",
    ANALYSIS / "posthoc_calibration_sensitivity" / "metrics.csv",
    ANALYSIS / "survival_offset" / "pooled_metrics.csv",
    ANALYSIS / "survival_offset" / "bootstrap_ci.csv",
    ANALYSIS / "survival_offset" / "harrell_c.csv",
    ANALYSIS / "survival_offset" / "permutation_importance.csv",
    ANALYSIS / "survival_descriptives" / "variables_by_domain.csv",
    ANALYSIS / "survival_descriptives" / "feature_missingness.csv",
    ANALYSIS / "survival_descriptives" / "event_frequency_by_month.csv",
]

ALLOWED_READINESS_STATUSES = {
    "Complete",
    "Author input required",
    "Manual check required",
}
ALLOWED_REFERENCE_STATUSES = {
    "Verified",
    "Author confirmation required",
}

MAIN_TEXT_WORD_LIMIT = 3000
ABSTRACT_WORD_LIMIT = 300
MAIN_EXHIBIT_LIMIT = 5
EXPECTED_ABSTRACT_HEADINGS = ["Background", "Methods", "Results", "Conclusions"]
WORD_RE = re.compile(r"\b[A-Za-z0-9]+(?:[-'][A-Za-z0-9]+)?\b")
CITATION_NUMBER_RE = re.compile(r"\d+")
PLACEHOLDER_RE = re.compile(r"\[[^\]]*(?:TBD|to be added)[^\]]*\]|\bto be added\b", re.I)


def _rel(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def count_page_breaks(path: Path) -> int:
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    return xml.count('<w:br w:type="page"/>') + xml.count('<w:br w:type="page"></w:br>')


def count_words(text: str) -> int:
    return len(WORD_RE.findall(text))


def section_text(doc: Document, start_heading: str, end_heading: str) -> str:
    collecting = False
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == start_heading:
            collecting = True
            continue
        if collecting and text == end_heading:
            break
        if collecting and text and not paragraph.style.name.startswith("Heading"):
            parts.append(text)
    return "\n".join(parts)


def section_headings(doc: Document, start_heading: str, end_heading: str) -> list[str]:
    collecting = False
    headings: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == start_heading:
            collecting = True
            continue
        if collecting and text == end_heading:
            break
        if collecting and text and paragraph.style.name.startswith("Heading"):
            headings.append(text)
    return headings


def heading_position(doc: Document, heading: str) -> int | None:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == heading:
            return index
    return None


def main_text_for_word_count(doc: Document) -> str:
    return section_text(doc, "Introduction", "References")


def clean_snippet(text: str, max_chars: int = 260) -> str:
    snippet = re.sub(r"\s+", " ", text).strip()
    if len(snippet) <= max_chars:
        return snippet
    return snippet[: max_chars - 3].rstrip() + "..."


def parse_citation_numbers(run: str) -> list[int]:
    numbers: list[int] = []
    for part in run.split(","):
        token = part.strip()
        if not token:
            continue
        if "-" in token:
            left, right = token.split("-", 1)
            if left.strip().isdigit() and right.strip().isdigit():
                start = int(left)
                end = int(right)
                if start <= end:
                    numbers.extend(range(start, end + 1))
                    continue
        numbers.extend(int(raw) for raw in CITATION_NUMBER_RE.findall(token))
    return numbers


def placeholder_audit_rows(path: Path, document_label: str, stop_heading: str | None = None) -> list[dict[str, str]]:
    doc = Document(path)
    rows: list[dict[str, str]] = []
    current_section = ""

    for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text.strip()
        if stop_heading and text == stop_heading:
            break
        if text and paragraph.style.name.startswith("Heading"):
            current_section = text
        for match in PLACEHOLDER_RE.finditer(paragraph.text):
            rows.append(
                {
                    "Document": document_label,
                    "Location": f"paragraph {paragraph_index}",
                    "Section": current_section,
                    "Placeholder": match.group(0),
                    "Context": clean_snippet(paragraph.text),
                }
            )

    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                for match in PLACEHOLDER_RE.finditer(cell.text):
                    rows.append(
                        {
                            "Document": document_label,
                            "Location": f"table {table_index}, row {row_index}, cell {cell_index}",
                            "Section": current_section,
                            "Placeholder": match.group(0),
                            "Context": clean_snippet(cell.text),
                        }
                    )
    return rows


def inspect_docx(path: Path, stop_heading: str | None = None) -> dict:
    doc = Document(path)
    paragraph_text: list[str] = []
    table_text: list[str] = []
    superscript_runs: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if stop_heading and text == stop_heading:
            break
        paragraph_text.append(paragraph.text)
        for run in paragraph.runs:
            if run.font.superscript:
                superscript_runs.append(run.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)

    joined = "\n".join(paragraph_text + table_text)
    placeholders = PLACEHOLDER_RE.findall(joined)
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": len(doc.inline_shapes),
        "exhibits": len(doc.tables) + len(doc.inline_shapes),
        "page_breaks": count_page_breaks(path),
        "superscript_runs": superscript_runs,
        "em_dash_count": joined.count("\u2014") + joined.count("\u2013"),
        "semicolon_count": joined.count(";"),
        "colon_count": joined.count(":"),
        "placeholder_count": len(placeholders),
        "placeholders": sorted(set(placeholders)),
    }


def citation_first_appearance(superscript_runs: list[str]) -> dict:
    seen: set[int] = set()
    all_numbers: list[int] = []
    first_order: list[int] = []
    out_of_order: list[dict[str, int | str]] = []
    expected_next = 1

    for run in superscript_runs:
        for number in parse_citation_numbers(run):
            all_numbers.append(number)
            if number in seen:
                continue
            if number != expected_next:
                out_of_order.append({"run": run, "number": number, "expected": expected_next})
            seen.add(number)
            first_order.append(number)
            expected_next = max(expected_next, number + 1)

    return {
        "all_numbers": all_numbers,
        "first_order": first_order,
        "out_of_order": out_of_order,
        "max_cited": max(seen) if seen else 0,
    }


def _read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def _validate_current_tables(errors: list[str], report: dict) -> None:
    onset = _read_csv(ANALYSIS / "survival_onset" / "pooled_metrics.csv")
    offset = _read_csv(ANALYSIS / "survival_offset" / "pooled_metrics.csv")
    table2 = _read_csv(TABLE_DIR / "table2_onset_performance.csv")
    table3 = _read_csv(TABLE_DIR / "table3_offset_performance.csv")
    table1 = _read_csv(TABLE_DIR / "table1_variables_by_domain.csv")
    expected_arches = {"logistic", "logistic_time_only", "lgbm", "lgbm_focal", "transformer"}

    report["onset_architectures"] = sorted(onset.get("architecture", pd.Series(dtype=str)).dropna().astype(str).tolist())
    report["offset_architectures"] = sorted(offset.get("architecture", pd.Series(dtype=str)).dropna().astype(str).tolist())
    report["table1_domain_rows"] = int(len(table1))
    report["table2_rows"] = int(len(table2))
    report["table3_rows"] = int(len(table3))

    if set(report["onset_architectures"]) != expected_arches:
        errors.append("Onset pooled metrics do not contain the five current architecture rows.")
    if set(report["offset_architectures"]) != expected_arches:
        errors.append("Offset pooled metrics do not contain the five current architecture rows.")
    if len(table2) != 5:
        errors.append("Current Table 2 should contain five onset architecture rows.")
    if len(table3) != 5:
        errors.append("Current Table 3 should contain five offset architecture rows.")
    if table1.empty or not {"domain", "n_variables"}.issubset(table1.columns):
        errors.append("Current Table 1 variables-by-domain table is missing required columns.")


def main() -> int:
    errors: list[str] = []
    advisories: list[str] = []
    report: dict = {}

    missing = [str(path) for path in REQUIRED_FILES if not path.exists() or path.stat().st_size == 0]
    report["missing_required_files"] = missing
    if missing:
        errors.append("Required active manuscript or survival artifact files are missing or empty.")

    if not MAIN_DOCX.exists() or not SUPPLEMENT_DOCX.exists():
        report["errors"] = errors
        report_path = OUT / "submission_preflight_report.json"
        report_path.write_text(json.dumps(report, indent=2))
        print(json.dumps(report, indent=2))
        return 1

    checklist = _read_csv(TABLE_DIR / "submission_readiness_checklist.csv")
    reporting = _read_csv(TABLE_DIR / "reporting_checklist_author_review.csv")
    reference_audit = _read_csv(TABLE_DIR / "reference_audit.csv")

    if not checklist.empty and "Status" in checklist:
        statuses = set(checklist["Status"].dropna().astype(str))
        report["readiness_status_counts"] = checklist["Status"].value_counts().to_dict()
        unknown_statuses = sorted(statuses - ALLOWED_READINESS_STATUSES)
        report["unknown_readiness_statuses"] = unknown_statuses
        if unknown_statuses:
            errors.append("Readiness checklist contains unknown statuses.")

    if not reporting.empty and "Status" in reporting:
        report["reporting_checklist_rows"] = len(reporting)
        report["reporting_checklist_status_counts"] = reporting["Status"].value_counts().to_dict()

    max_reference = 0
    if not reference_audit.empty:
        report["reference_audit_rows"] = len(reference_audit)
        if "Verification status" in reference_audit:
            reference_statuses = set(reference_audit["Verification status"].dropna().astype(str))
            report["reference_audit_status_counts"] = reference_audit["Verification status"].value_counts().to_dict()
            unknown_reference_statuses = sorted(reference_statuses - ALLOWED_REFERENCE_STATUSES)
            report["unknown_reference_statuses"] = unknown_reference_statuses
            if unknown_reference_statuses:
                errors.append("Reference audit contains unknown verification statuses.")
            if "Author confirmation required" in reference_statuses:
                advisories.append("Reference audit contains sources requiring author or license-holder confirmation.")
        if "Number" in reference_audit:
            reference_numbers = reference_audit["Number"].astype(int).tolist()
            expected_reference_numbers = list(range(1, len(reference_audit) + 1))
            report["reference_numbers_sequential"] = reference_numbers == expected_reference_numbers
            max_reference = max(reference_numbers) if reference_numbers else 0
            if reference_numbers != expected_reference_numbers:
                errors.append("Reference audit numbers are not sequential.")

    main_doc = inspect_docx(MAIN_DOCX, stop_heading="References")
    supplement_doc = inspect_docx(SUPPLEMENT_DOCX)
    report["main_docx"] = main_doc
    report["supplement_docx"] = supplement_doc

    placeholder_context_rows = (
        placeholder_audit_rows(MAIN_DOCX, MAIN_DOCX.name, stop_heading="References")
        + placeholder_audit_rows(SUPPLEMENT_DOCX, SUPPLEMENT_DOCX.name)
    )
    placeholder_audit_path = TABLE_DIR / "submission_placeholder_audit.csv"
    pd.DataFrame(
        placeholder_context_rows,
        columns=["Document", "Location", "Section", "Placeholder", "Context"],
    ).to_csv(placeholder_audit_path, index=False)
    report["placeholder_audit_file"] = _rel(placeholder_audit_path)
    report["author_input_placeholder_count"] = len(placeholder_context_rows)
    report["placeholder_contexts"] = placeholder_context_rows
    if placeholder_context_rows:
        advisories.append("Active Word documents contain unresolved author-input placeholders.")

    document = Document(MAIN_DOCX)
    abstract_text = section_text(document, "Abstract", "Introduction")
    abstract_words = count_words(abstract_text)
    abstract_headings = section_headings(document, "Abstract", "Introduction")
    if not abstract_headings:
        positions = []
        for heading in EXPECTED_ABSTRACT_HEADINGS:
            match = re.search(rf"\b{re.escape(heading)}\.", abstract_text)
            positions.append(match.start() if match else -1)
        if all(pos >= 0 for pos in positions) and positions == sorted(positions):
            abstract_headings = EXPECTED_ABSTRACT_HEADINGS.copy()
    main_text_words = count_words(main_text_for_word_count(document))
    references_position = heading_position(document, "References")
    exhibits_position = heading_position(document, "Tables and figures")

    report["abstract_word_count"] = abstract_words
    report["abstract_word_limit"] = ABSTRACT_WORD_LIMIT
    report["abstract_headings"] = abstract_headings
    report["expected_abstract_headings"] = EXPECTED_ABSTRACT_HEADINGS
    report["main_text_word_count"] = main_text_words
    report["main_text_word_limit"] = MAIN_TEXT_WORD_LIMIT
    report["main_exhibit_limit"] = MAIN_EXHIBIT_LIMIT
    report["references_heading_position"] = references_position
    report["tables_figures_heading_position"] = exhibits_position

    if main_doc["em_dash_count"] or main_doc["semicolon_count"] or main_doc["colon_count"]:
        errors.append("Main manuscript contains disallowed punctuation before references or in exhibit text.")
    if supplement_doc["em_dash_count"] or supplement_doc["semicolon_count"] or supplement_doc["colon_count"]:
        errors.append("Supplement contains disallowed punctuation.")
    if main_doc["exhibits"] > MAIN_EXHIBIT_LIMIT:
        errors.append("Main manuscript exceeds the local five-exhibit limit.")
    if abstract_words > ABSTRACT_WORD_LIMIT:
        errors.append("Abstract exceeds the local 300-word limit.")
    if abstract_headings != EXPECTED_ABSTRACT_HEADINGS:
        errors.append("Structured abstract headings do not match the expected structure.")
    if main_text_words > MAIN_TEXT_WORD_LIMIT:
        errors.append("Main text exceeds the local 3000-word limit.")
    if references_position is None:
        errors.append("References heading is missing.")
    if exhibits_position is None:
        advisories.append("No literal Tables and figures heading was found, table and figure titles appear after References.")
    elif exhibits_position <= references_position:
        errors.append("Tables and figures section must appear after References.")

    citation_order = citation_first_appearance(main_doc["superscript_runs"])
    report["main_citation_first_appearance"] = citation_order
    if citation_order["out_of_order"]:
        errors.append("Main manuscript citations are not numbered in order of first appearance.")
    if max_reference and citation_order["max_cited"] > max_reference:
        errors.append("Main manuscript cites a reference number not present in the reference audit.")

    _validate_current_tables(errors, report)

    report["advisories"] = advisories
    report["errors"] = errors
    report_path = OUT / "submission_preflight_report.json"
    report_path.write_text(json.dumps(report, indent=2))
    print(json.dumps(report, indent=2))
    return 1 if errors else 0


if __name__ == "__main__":
    sys.exit(main())
